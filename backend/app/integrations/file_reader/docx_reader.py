from collections.abc import Iterator
from typing import Any

from app.integrations.file_reader.base import FileSource
from app.integrations.file_reader.docx_table import table_text_and_metadata
from app.integrations.file_reader.errors import FileReaderDependencyError, FileReaderError
from app.integrations.file_reader.structure import DocumentStructureAnalyzer
from app.integrations.file_reader.structure.schema import (
    ParsedBlock,
    ParsedDocumentV1,
    ParsedMetadata,
    SourceLocation,
    StyleFeatures,
)


class DocxReader:
    """DOCX 读取器。

    Input:
        ``FileSource``，其 ``path`` 指向本地 ``.docx`` 文件。

    Output:
        ``ParsedDocumentV1``。其中 ``blocks`` 保存按文档顺序读取的段落和表格，
        ``sections`` 由 ``DocumentStructureAnalyzer`` 在 blocks 基础上推断。

    Processing:
        1. 校验路径和扩展名。
        2. 用 python-docx 打开 Word 文件。
        3. 按 OOXML body 顺序遍历段落和表格，生成 ``ParsedBlock``。
        4. 抽取段落样式、对齐、缩进、粗体比例和显式字号等版式特征。
        5. 调用章节推断器补充 ``sections`` 和 ``warnings``。

    Boundary:
        本类只读取 Word 文件里的事实信息；合同章节层级不能直接依赖 Word Heading 样式，
        必须由结构推断层结合文本编号和版式特征判断。
    """

    reader_type = "python-docx"

    async def parse(self, source: FileSource) -> ParsedDocumentV1:
        """读取 .docx 文件并返回 ParsedDocument v1。

        Args:
            source: 文件来源，``source.file_type`` 必须为 ``docx``。

        Returns:
            统一解析结果。空段落会被跳过；表格会作为一个 ``kind="table"``
            的 block 输出，并在 metadata 中保留原始单元格网格。

        Raises:
            FileReaderError: 文件不存在或扩展名不是 docx。
            FileReaderDependencyError: 未安装 python-docx。
        """
        if not source.path.exists():
            raise FileReaderError(f"file not found: {source.path}")
        if source.file_type != "docx":
            raise FileReaderError("DocxReader only supports .docx files")

        try:
            from docx import Document
        except ImportError as exc:
            raise FileReaderDependencyError("python-docx is required for .docx parsing") from exc

        document = Document(str(source.path))
        blocks: list[ParsedBlock] = []
        paragraph_count = 0
        table_count = 0

        for order, item in enumerate[tuple[str, Any]](_iter_document_items(document), start=1):
            item_type, value = item
            if item_type == "paragraph":
                text = " ".join((value.text or "").split())
                if not text:
                    continue
                paragraph_count += 1
                blocks.append(
                    ParsedBlock(
                        id=f"b-{len(blocks) + 1:06d}",
                        kind="paragraph",
                        text=text,
                        order=order,
                        source_location=SourceLocation(paragraph_index=paragraph_count),
                        style_features=_paragraph_style_features(value),
                    )
                )
            elif item_type == "table":
                table_count += 1
                text, metadata = table_text_and_metadata(value)
                if not text:
                    continue
                blocks.append(
                    ParsedBlock(
                        id=f"b-{len(blocks) + 1:06d}",
                        kind="table",
                        text=text,
                        order=order,
                        source_location=SourceLocation(table_index=table_count),
                        style_features=StyleFeatures(style_name="Table"),
                        metadata=metadata,
                    )
                )

        parsed = ParsedDocumentV1(
            metadata=ParsedMetadata(
                filename=source.filename,
                file_type=source.file_type,
                reader_type=self.reader_type,
                paragraph_count=paragraph_count,
                table_count=table_count,
            ),
            blocks=blocks,
        )
        return DocumentStructureAnalyzer().analyze(parsed)


def _iter_document_items(document: Any) -> Iterator[tuple[str, Any]]:
    """按 Word 文档真实顺序遍历段落和表格。

    Args:
        document: python-docx 的 ``Document`` 对象。

    Yields:
        ``("paragraph", Paragraph)`` 或 ``("table", Table)``。

    Processing:
        直接遍历 OOXML ``document.element.body``，遇到 ``CT_P`` 包装成 Paragraph，
        遇到 ``CT_Tbl`` 包装成 Table。

    Reason:
        python-docx 的 ``document.paragraphs`` 与 ``document.tables`` 会分别列出内容，
        无法保留表格在正文中的位置；直接遍历 OOXML body 可以保证 blocks 顺序稳定。
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P

    body = document.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield "paragraph", Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield "table", Table(child, document)


def _paragraph_style_features(paragraph: Any) -> StyleFeatures:
    """抽取章节推断需要的段落版式信号。

    Args:
        paragraph: python-docx 的 ``Paragraph`` 对象。

    Returns:
        ``StyleFeatures``，包含样式名、对齐方式、首行缩进、左缩进、
        粗体字符占比和最常见显式字号。

    Processing:
        1. 从 paragraph_format 读取缩进和对齐。
        2. 遍历 runs，统计粗体字符数和显式字号。
        3. 将 python-docx 的枚举值归一化为平台内部字符串。

    Reason:
        很多合同全篇使用 Normal 样式，所以这些特征只作为辅助信号，
        不能替代文本编号和上下文规则。
    """
    fmt = paragraph.paragraph_format
    bold_chars = 0
    sized_runs: list[float] = []
    total_chars = 0
    for run in paragraph.runs:
        text_len = len(run.text or "")
        total_chars += text_len
        if run.bold:
            bold_chars += text_len
        if run.font.size is not None:
            sized_runs.append(run.font.size.pt)

    return StyleFeatures(
        style_name=paragraph.style.name if paragraph.style else None,
        alignment=_alignment_name(paragraph.alignment),
        first_line_indent=int(fmt.first_line_indent) if fmt.first_line_indent is not None else None,
        left_indent=int(fmt.left_indent) if fmt.left_indent is not None else None,
        bold_ratio=(bold_chars / total_chars) if total_chars else 0.0,
        font_size_pt=_most_common_size(sized_runs),
    )


def _alignment_name(value: Any) -> str:
    """把 python-docx 的对齐枚举归一化为平台内部字符串。

    Args:
        value: python-docx 的段落对齐枚举，可能为 None。

    Returns:
        ``left``、``center``、``right``、``justify`` 之一；未知值保守归为 ``left``。
    """
    if value is None:
        return "left"
    name = getattr(value, "name", str(value)).lower()
    if "center" in name:
        return "center"
    if "right" in name:
        return "right"
    if "justify" in name:
        return "justify"
    return "left"


def _most_common_size(values: list[float]) -> float | None:
    """返回段落 run 中最常见的显式字号。

    Args:
        values: 从 runs 中抽取到的字号 pt 列表。

    Returns:
        出现次数最多的字号；如果段落没有显式字号，则返回 None。
    """
    if not values:
        return None
    return max(set[float](values), key=values.count)
